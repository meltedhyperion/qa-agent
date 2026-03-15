"""Parse .docx test documents and extract sections + test cases."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph


def parse_document(file_path: str) -> dict:
    """Parse a .docx file and return structured sections with test cases."""
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"Document not found: {file_path}")

    doc = Document(str(path))
    sections = _extract_sections(doc)

    total_cases = sum(len(s["test_cases"]) for s in sections)
    title = _extract_document_title(doc)

    return {
        "filename": path.name,
        "sections": sections,
        "total_sections": len(sections),
        "total_test_cases": total_cases,
        "document_title": title,
    }


def list_sections(file_path: str) -> dict:
    """Quick scan returning section names and test case counts."""
    result = parse_document(file_path)
    return {
        "sections": [
            {"name": s["name"], "test_case_count": len(s["test_cases"])}
            for s in result["sections"]
        ]
    }


def extract_test_cases(file_path: str, section_name: str) -> dict:
    """Extract test cases from a specific section."""
    result = parse_document(file_path)
    for section in result["sections"]:
        if section["name"].strip().lower() == section_name.strip().lower():
            return {"section": section["name"], "test_cases": section["test_cases"]}
    return {"section": section_name, "test_cases": [], "error": "Section not found"}


# ── Internal helpers ─────────────────────────────────────────────────────────


def _extract_document_title(doc: Document) -> str:
    """Try to extract a document title from the first heading or Title style."""
    for para in doc.paragraphs[:10]:
        if para.style and para.style.name and (
            "Title" in para.style.name or "Heading" in para.style.name
        ):
            text = para.text.strip()
            if text:
                return text
    return ""


def _extract_sections(doc: Document) -> list[dict]:
    """Walk through the document, splitting by headings, and extract test cases
    from tables that follow each heading."""
    sections: list[dict] = []

    # Gather all body elements in document order (paragraphs + tables)
    body_elements = _iter_body_elements(doc)

    current_section: dict | None = None
    tc_counter = 0

    for element in body_elements:
        if isinstance(element, Paragraph):
            heading_level = _get_heading_level(element)
            if heading_level is not None and element.text.strip():
                # Start a new section
                current_section = {
                    "name": element.text.strip(),
                    "heading_level": heading_level,
                    "test_cases": [],
                }
                sections.append(current_section)

        elif isinstance(element, Table) and current_section is not None:
            cases = _extract_test_cases_from_table(element, tc_counter)
            current_section["test_cases"].extend(cases)
            tc_counter += len(cases)

    # If no heading-based sections found, treat the entire doc as one section
    if not sections:
        all_cases = []
        for element in body_elements:
            if isinstance(element, Table):
                cases = _extract_test_cases_from_table(element, tc_counter)
                all_cases.extend(cases)
                tc_counter += len(cases)
        if all_cases:
            sections.append({
                "name": "All Tests",
                "heading_level": 1,
                "test_cases": all_cases,
            })

    return sections


def _iter_body_elements(doc: Document):
    """Iterate paragraphs and tables in document body order."""
    from docx.oxml.ns import qn

    body = doc.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, doc)
        elif child.tag == qn("w:tbl"):
            yield Table(child, doc)


def _get_heading_level(paragraph: Paragraph) -> int | None:
    """Return heading level (1-9) or None if not a heading."""
    style = paragraph.style
    if style is None:
        return None
    name = style.name or ""
    if name == "Title":
        return 0
    match = re.match(r"Heading\s*(\d)", name)
    if match:
        return int(match.group(1))
    return None


def _extract_test_cases_from_table(table: Table, start_id: int = 0) -> list[dict]:
    """Extract test cases from a table.

    Handles common formats:
      - 2-col: Title | Description/Steps
      - 3-col: ID | Title | Steps
      - 4-col: ID | Title | Steps | Expected Result
      - Also detects header row and skips it.
    """
    rows = table.rows
    if len(rows) < 2:
        return []

    num_cols = len(table.columns)
    if num_cols < 2:
        return []

    # Detect header row
    first_row_texts = [_cell_text(c) for c in rows[0].cells]
    has_header = _looks_like_header(first_row_texts)
    data_rows = rows[1:] if has_header else rows

    # Determine column mapping
    col_map = _detect_column_mapping(first_row_texts, num_cols)

    test_cases = []
    for i, row in enumerate(data_rows):
        cells = row.cells
        if len(cells) < 2:
            continue

        tc_id = ""
        title = ""
        steps_text = ""
        expected = ""

        if col_map["id_col"] is not None and col_map["id_col"] < len(cells):
            tc_id = _cell_text(cells[col_map["id_col"]]).strip()
        if col_map["title_col"] is not None and col_map["title_col"] < len(cells):
            title = _cell_text(cells[col_map["title_col"]]).strip()
        if col_map["steps_col"] is not None and col_map["steps_col"] < len(cells):
            steps_text = _cell_text(cells[col_map["steps_col"]]).strip()
        if col_map["expected_col"] is not None and col_map["expected_col"] < len(cells):
            expected = _cell_text(cells[col_map["expected_col"]]).strip()

        # Skip empty rows
        if not title and not steps_text:
            continue

        # If no separate title column, use the first non-empty cell
        if not title and steps_text:
            title = steps_text[:80]

        # Assign an ID if not present
        if not tc_id:
            tc_id = f"TC-{start_id + len(test_cases) + 1:03d}"

        # Parse steps from text
        steps = _parse_steps(steps_text)

        test_cases.append({
            "id": tc_id,
            "title": title,
            "description": steps_text,
            "steps": steps,
            "expected_result": expected,
            "raw_text": steps_text,
        })

    return test_cases


def _cell_text(cell) -> str:
    """Get all text from a table cell, preserving paragraph breaks."""
    return "\n".join(p.text for p in cell.paragraphs)


def _looks_like_header(texts: list[str]) -> bool:
    """Heuristic: a row is a header if it contains typical header keywords."""
    header_keywords = {
        "test", "case", "id", "title", "name", "description",
        "steps", "step", "scenario", "expected", "result", "action",
        "s.no", "sr", "sl", "#", "no.", "s no",
    }
    combined = " ".join(t.lower() for t in texts)
    matches = sum(1 for kw in header_keywords if kw in combined)
    return matches >= 2


def _detect_column_mapping(header_texts: list[str], num_cols: int) -> dict:
    """Figure out which column is which based on header text or column count."""
    mapping = {"id_col": None, "title_col": None, "steps_col": None, "expected_col": None}

    # Try to match by header text
    lower_headers = [t.lower().strip() for t in header_texts]
    for i, h in enumerate(lower_headers):
        if any(kw in h for kw in ("id", "s.no", "sr", "sl", "#", "no.")):
            mapping["id_col"] = i
        elif any(kw in h for kw in ("title", "name", "scenario", "test case")):
            mapping["title_col"] = i
        elif any(kw in h for kw in ("step", "description", "action", "procedure")):
            mapping["steps_col"] = i
        elif any(kw in h for kw in ("expected", "result", "output")):
            mapping["expected_col"] = i

    # If header detection worked, return
    if mapping["title_col"] is not None:
        # If steps_col not found but we have 3+ cols, use the column after title
        if mapping["steps_col"] is None and num_cols > 2:
            for i in range(num_cols):
                if i not in (mapping["id_col"], mapping["title_col"], mapping["expected_col"]):
                    mapping["steps_col"] = i
                    break
        return mapping

    # Fallback: positional mapping based on column count
    if num_cols == 2:
        mapping["title_col"] = 0
        mapping["steps_col"] = 1
    elif num_cols == 3:
        mapping["id_col"] = 0
        mapping["title_col"] = 1
        mapping["steps_col"] = 2
    elif num_cols >= 4:
        mapping["id_col"] = 0
        mapping["title_col"] = 1
        mapping["steps_col"] = 2
        mapping["expected_col"] = 3

    return mapping


def _parse_steps(text: str) -> list[str]:
    """Parse numbered or bulleted steps from text.

    Handles:
      1. Step one
      2. Step two
      - Step three
      * Step four
      Step five (plain lines)
    """
    if not text.strip():
        return []

    lines = text.strip().split("\n")
    steps: list[str] = []

    # Try to detect numbered pattern
    numbered_pattern = re.compile(r"^\s*(\d+)[.):\-]\s*(.+)")
    bullet_pattern = re.compile(r"^\s*[-•*]\s*(.+)")

    for line in lines:
        line = line.strip()
        if not line:
            continue

        m = numbered_pattern.match(line)
        if m:
            steps.append(m.group(2).strip())
            continue

        m = bullet_pattern.match(line)
        if m:
            steps.append(m.group(1).strip())
            continue

        # Plain line — could be a continuation or a standalone step
        if steps and len(line) < 20 and not line[0].isupper():
            # Likely a continuation of the previous step
            steps[-1] += " " + line
        else:
            steps.append(line)

    return steps

"""Generate formatted .docx result reports with embedded screenshots."""

from __future__ import annotations

import os
from datetime import datetime
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from PIL import Image


def generate_report(
    session_path: str,
    test_results: list[dict],
    tester_name: str = "GPT-4o",
    test_date: str | None = None,
) -> dict:
    """Generate a formatted results .docx with a 4-column table.

    Args:
        session_path: Path to session export folder.
        test_results: List of dicts, each containing:
            - test_title: str
            - steps: list[str]
            - screenshot_paths: list[str]
            - status: "passed" | "failed" | "skipped"
        tester_name: Name for the "Verified By" column.
        test_date: Date string (defaults to today).

    Returns:
        {"report_path": "...", "success": True}
    """
    if test_date is None:
        test_date = datetime.now().strftime("%Y-%m-%d")

    doc = Document()

    # ── Page setup: landscape, narrow margins ────────────────────────────
    for section in doc.sections:
        section.orientation = WD_ORIENT.LANDSCAPE
        # Swap width/height for landscape
        section.page_width, section.page_height = section.page_height, section.page_width
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    # ── Title ────────────────────────────────────────────────────────────
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("Test Execution Report")
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = "Calibri"

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run(f"Generated: {test_date}  |  Tester: {tester_name}")
    sub_run.font.size = Pt(10)
    sub_run.font.color.rgb = RGBColor(100, 100, 100)
    sub_run.font.name = "Calibri"

    doc.add_paragraph()  # spacer

    # ── Summary ──────────────────────────────────────────────────────────
    passed = sum(1 for r in test_results if r.get("status") == "passed")
    failed = sum(1 for r in test_results if r.get("status") == "failed")
    skipped = sum(1 for r in test_results if r.get("status") == "skipped")
    total = len(test_results)

    summary = doc.add_paragraph()
    summary.alignment = WD_ALIGN_PARAGRAPH.LEFT
    s_run = summary.add_run(
        f"Total: {total}  |  Passed: {passed}  |  Failed: {failed}  |  Skipped: {skipped}"
    )
    s_run.font.size = Pt(10)
    s_run.font.name = "Calibri"

    doc.add_paragraph()  # spacer

    # ── Results table ────────────────────────────────────────────────────
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    header_cells = table.rows[0].cells
    headers = ["Test Title", "Steps", "Screenshots", "Verified By"]
    for i, header_text in enumerate(headers):
        p = header_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header_text)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = "Calibri"
        # Header background
        _set_cell_shading(header_cells[i], "2F5496")

        run.font.color.rgb = RGBColor(255, 255, 255)

    # Set header row to repeat on each page
    _set_repeat_header_row(table)

    # Column widths (landscape ≈ 10" usable)
    col_widths = [Inches(2.0), Inches(2.5), Inches(4.0), Inches(1.5)]

    # ── Data rows ────────────────────────────────────────────────────────
    for result in test_results:
        row = table.add_row()
        cells = row.cells

        # Col 0: Test Title
        _write_cell_text(
            cells[0],
            result.get("test_title", "Untitled"),
            bold=True,
            size=Pt(10),
        )

        # Col 1: Steps (numbered)
        steps = result.get("steps", [])
        steps_text = "\n".join(f"{i + 1}. {step}" for i, step in enumerate(steps))
        _write_cell_text(cells[1], steps_text, size=Pt(9))

        # Col 2: Screenshots (stacked vertically)
        screenshot_paths = result.get("screenshot_paths", [])
        _add_screenshots_to_cell(cells[2], screenshot_paths)

        # Col 3: Verified By
        verified_text = f"{tester_name}\n{test_date}"
        status = result.get("status", "unknown")
        status_text = f"\nStatus: {status.upper()}"
        _write_cell_text(cells[3], verified_text + status_text, size=Pt(8))

    # Apply column widths
    for row in table.rows:
        for i, width in enumerate(col_widths):
            if i < len(row.cells):
                row.cells[i].width = width

    # ── Save ─────────────────────────────────────────────────────────────
    report_path = os.path.join(session_path, "results_report.docx")
    Path(session_path).mkdir(parents=True, exist_ok=True)
    doc.save(report_path)

    return {"report_path": report_path, "success": True}


def append_test_result(
    report_path: str,
    test_result: dict,
    tester_name: str = "GPT-4o",
    test_date: str | None = None,
) -> dict:
    """Add a single test result row to an existing report. Creates the report
    if it doesn't exist."""
    if test_date is None:
        test_date = datetime.now().strftime("%Y-%m-%d")

    if not os.path.exists(report_path):
        # Create new report with just this result
        session_path = os.path.dirname(report_path)
        return generate_report(session_path, [test_result], tester_name, test_date)

    doc = Document(report_path)

    # Find the table (should be the last one)
    if not doc.tables:
        return {"success": False, "error": "No table found in existing report"}

    table = doc.tables[-1]
    row = table.add_row()
    cells = row.cells

    _write_cell_text(
        cells[0],
        test_result.get("test_title", "Untitled"),
        bold=True,
        size=Pt(10),
    )

    steps = test_result.get("steps", [])
    steps_text = "\n".join(f"{i + 1}. {step}" for i, step in enumerate(steps))
    _write_cell_text(cells[1], steps_text, size=Pt(9))

    _add_screenshots_to_cell(cells[2], test_result.get("screenshot_paths", []))

    status = test_result.get("status", "unknown")
    verified_text = f"{tester_name}\n{test_date}\nStatus: {status.upper()}"
    _write_cell_text(cells[3], verified_text, size=Pt(8))

    doc.save(report_path)
    return {"report_path": report_path, "success": True}


# ── Helpers ──────────────────────────────────────────────────────────────────


def _write_cell_text(
    cell,
    text: str,
    bold: bool = False,
    size=Pt(10),
):
    """Write text into a cell, clearing existing content."""
    # Clear default empty paragraph
    cell.paragraphs[0].clear()
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = size
    run.font.name = "Calibri"


def _add_screenshots_to_cell(cell, screenshot_paths: list[str]):
    """Add screenshot images stacked vertically inside a cell."""
    cell.paragraphs[0].clear()

    for i, img_path in enumerate(screenshot_paths):
        if not os.path.exists(img_path):
            continue

        # Resize if too large and get dimensions
        width = _get_image_width(img_path)

        if i > 0:
            # Add spacing paragraph between images
            cell.add_paragraph()

        para = cell.paragraphs[-1] if i > 0 else cell.paragraphs[0]
        run = para.add_run()
        run.add_picture(img_path, width=width)


def _get_image_width(img_path: str, max_width_inches: float = 3.8) -> Inches:
    """Calculate appropriate image width for the cell."""
    try:
        with Image.open(img_path) as img:
            w, h = img.size
            aspect = h / w if w > 0 else 1
            # If very wide image, use full width; if tall, shrink
            if aspect > 1.5:
                return Inches(max_width_inches * 0.6)
            return Inches(max_width_inches)
    except Exception:
        return Inches(max_width_inches)


def _set_cell_shading(cell, color_hex: str):
    """Set cell background color."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def _set_repeat_header_row(table):
    """Make the first row repeat as header on each page."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tr = table.rows[0]._tr
    trPr = tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    trPr.append(tbl_header)

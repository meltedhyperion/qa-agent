"""
Generate sample .docx test fixture files for the QA agent.

Creates three documents with different formats to test the rigid parser
and LLM parser:
  1. sample_test_doc.docx              — Standard 4-column table format
  2. sample_test_doc_paragraph_steps.docx — Steps written as paragraphs (not in tables)
  3. sample_test_doc_unstructured.docx — Free-form prose, no tables or numbered lists

Usage:
    python tests/fixtures/generate_test_docs.py
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

FIXTURES_DIR = Path(__file__).parent


def set_cell_text(cell, text, bold=False):
    """Set cell text with optional bold formatting."""
    cell.text = ""
    run = cell.paragraphs[0].add_run(text)
    run.bold = bold
    run.font.size = Pt(10)


def add_metadata(doc, lines):
    """Add metadata paragraphs below the title."""
    for line in lines:
        p = doc.add_paragraph(line)
        p.style = doc.styles["Normal"]


# ---------------------------------------------------------------------------
# 1. sample_test_doc.docx — Standard table-based format
# ---------------------------------------------------------------------------
def generate_sample_test_doc():
    doc = Document()

    # Title
    title = doc.add_paragraph("SauceDemo Application Test Suite")
    title.style = doc.styles["Title"]

    # Metadata
    add_metadata(doc, [
        "Application URL: https://www.saucedemo.com",
        "Version: 1.0",
        "Date: 2026-03-15",
        "Environment: Production",
    ])

    # --- Section 1: Smoke Tests ---
    doc.add_heading("Smoke Tests", level=1)
    doc.add_paragraph(
        "Basic smoke tests to verify core application functionality."
    )

    smoke_tests = [
        ["Test ID", "Test Title", "Test Steps", "Expected Result"],
        [
            "TC-001",
            "Login with valid credentials",
            (
                "1. Navigate to https://www.saucedemo.com\n"
                "2. Enter 'standard_user' in the Username field\n"
                "3. Enter 'secret_sauce' in the Password field\n"
                "4. Click the 'Login' button"
            ),
            "User is logged in and the inventory/products page is displayed.",
        ],
        [
            "TC-002",
            "Login with invalid credentials",
            (
                "1. Navigate to https://www.saucedemo.com\n"
                "2. Enter 'invalid_user' in the Username field\n"
                "3. Enter 'wrong_password' in the Password field\n"
                "4. Click the 'Login' button"
            ),
            (
                "An error message is displayed: "
                "'Username and password do not match any user in this service.'"
            ),
        ],
        [
            "TC-003",
            "Login with locked-out user",
            (
                "1. Navigate to https://www.saucedemo.com\n"
                "2. Enter 'locked_out_user' in the Username field\n"
                "3. Enter 'secret_sauce' in the Password field\n"
                "4. Click the 'Login' button"
            ),
            "An error message is displayed: 'Sorry, this user has been locked out.'",
        ],
    ]

    table = doc.add_table(rows=len(smoke_tests), cols=4)
    table.style = "Table Grid"
    for i, row_data in enumerate(smoke_tests):
        for j, cell_text in enumerate(row_data):
            set_cell_text(table.rows[i].cells[j], cell_text, bold=(i == 0))

    # --- Section 2: Functional Tests ---
    doc.add_heading("Functional Tests", level=1)
    doc.add_paragraph(
        "Functional tests covering product browsing, cart operations, and checkout flow."
    )

    functional_tests = [
        ["Test ID", "Test Title", "Test Steps", "Expected Result"],
        [
            "TC-004",
            "View product inventory",
            (
                "1. Login with valid credentials (standard_user / secret_sauce)\n"
                "2. Verify the products page is displayed\n"
                "3. Verify that at least one product is listed with a name, price, and 'Add to cart' button"
            ),
            "Products page shows a list of products with names, prices, and add-to-cart buttons.",
        ],
        [
            "TC-005",
            "Add product to cart",
            (
                "1. Login with valid credentials\n"
                "2. Click 'Add to cart' on the first product\n"
                "3. Verify the cart badge shows '1'"
            ),
            "Cart icon badge displays '1' indicating one item added.",
        ],
        [
            "TC-006",
            "Remove product from cart",
            (
                "1. Login with valid credentials\n"
                "2. Add a product to the cart\n"
                "3. Click 'Remove' on the same product\n"
                "4. Verify the cart badge is no longer visible or shows '0'"
            ),
            "Product is removed and cart badge disappears or resets.",
        ],
        [
            "TC-007",
            "View cart contents",
            (
                "1. Login with valid credentials\n"
                "2. Add 'Sauce Labs Backpack' to the cart\n"
                "3. Click the cart icon\n"
                "4. Verify the cart page lists 'Sauce Labs Backpack'"
            ),
            "Cart page displays the added product with its name and price.",
        ],
        [
            "TC-008",
            "Sort products by price low to high",
            (
                "1. Login with valid credentials\n"
                "2. Click the sort dropdown on the products page\n"
                "3. Select 'Price (low to high)'\n"
                "4. Verify products are sorted by ascending price"
            ),
            "Products are reordered with the cheapest item first.",
        ],
        [
            "TC-009",
            "Complete checkout",
            (
                "1. Login with valid credentials\n"
                "2. Add a product to the cart\n"
                "3. Click the cart icon\n"
                "4. Click 'Checkout'\n"
                "5. Fill in first name, last name, and postal code\n"
                "6. Click 'Continue'\n"
                "7. Click 'Finish'"
            ),
            "Order confirmation page is displayed with 'Thank you for your order!' message.",
        ],
        [
            "TC-010",
            "Logout",
            (
                "1. Login with valid credentials\n"
                "2. Click the hamburger menu icon\n"
                "3. Click 'Logout'\n"
                "4. Verify the login page is displayed"
            ),
            "User is redirected back to the login page.",
        ],
        [
            "TC-011",
            "Reset app state",
            (
                "1. Login with valid credentials\n"
                "2. Add two products to the cart\n"
                "3. Open the hamburger menu\n"
                "4. Click 'Reset App State'\n"
                "5. Verify the cart badge is removed"
            ),
            "Cart is cleared and all 'Remove' buttons revert to 'Add to cart'.",
        ],
    ]

    table2 = doc.add_table(rows=len(functional_tests), cols=4)
    table2.style = "Table Grid"
    for i, row_data in enumerate(functional_tests):
        for j, cell_text in enumerate(row_data):
            set_cell_text(table2.rows[i].cells[j], cell_text, bold=(i == 0))

    path = FIXTURES_DIR / "sample_test_doc.docx"
    doc.save(str(path))
    print(f"Created {path}")


# ---------------------------------------------------------------------------
# 2. sample_test_doc_paragraph_steps.docx — Paragraph-style steps
# ---------------------------------------------------------------------------
def generate_paragraph_steps_doc():
    doc = Document()

    title = doc.add_paragraph("SauceDemo Test Suite — Paragraph Format")
    title.style = doc.styles["Title"]

    add_metadata(doc, [
        "Application: https://www.saucedemo.com",
        "Author: QA Team",
        "Date: 2026-03-15",
    ])

    # --- Section 1: Authentication Tests ---
    doc.add_heading("Authentication Tests", level=1)
    doc.add_paragraph(
        "These tests verify login and authentication behavior."
    )

    # TC-001
    doc.add_heading("TC-001: Successful Login", level=2)
    doc.add_paragraph(
        "Navigate to the SauceDemo login page at https://www.saucedemo.com. "
        "Enter the username 'standard_user' and the password 'secret_sauce'. "
        "Click the Login button. The user should be redirected to the products page "
        "showing a list of inventory items."
    )
    doc.add_paragraph("Expected Result: The products/inventory page loads successfully.")

    # TC-002
    doc.add_heading("TC-002: Failed Login with Bad Password", level=2)
    doc.add_paragraph(
        "Open the SauceDemo login page. Type 'standard_user' as the username "
        "and 'bad_password' as the password. Press the Login button. "
        "An error banner should appear with the message "
        "'Username and password do not match any user in this service.'"
    )
    doc.add_paragraph(
        "Expected Result: Error message displayed, user remains on login page."
    )

    # TC-003
    doc.add_heading("TC-003: Locked-Out User", level=2)
    doc.add_paragraph(
        "Go to the login page. Enter 'locked_out_user' for the username and "
        "'secret_sauce' for the password. Click Login. The application should "
        "display a message saying 'Sorry, this user has been locked out.'"
    )
    doc.add_paragraph(
        "Expected Result: Locked-out error is shown; login is denied."
    )

    # --- Section 2: Shopping Cart Tests ---
    doc.add_heading("Shopping Cart Tests", level=1)
    doc.add_paragraph(
        "These tests cover adding, removing, and viewing items in the cart."
    )

    # TC-004
    doc.add_heading("TC-004: Add Item to Cart", level=2)
    doc.add_paragraph(
        "Log in with standard_user / secret_sauce. On the products page, "
        "find the 'Sauce Labs Backpack' and click its 'Add to cart' button. "
        "The cart badge in the top-right corner should update to show '1'. "
        "The button text should change from 'Add to cart' to 'Remove'."
    )
    doc.add_paragraph("Expected Result: Cart badge shows 1; button says Remove.")

    # TC-005
    doc.add_heading("TC-005: Remove Item from Cart", level=2)
    doc.add_paragraph(
        "After adding the Sauce Labs Backpack to the cart, click the 'Remove' "
        "button next to it. The cart badge should disappear or show 0 items. "
        "The button should revert to 'Add to cart'."
    )
    doc.add_paragraph("Expected Result: Cart badge gone; button reverts to Add to cart.")

    # TC-006
    doc.add_heading("TC-006: View Cart Page", level=2)
    doc.add_paragraph(
        "Log in and add two items to the cart: 'Sauce Labs Backpack' and "
        "'Sauce Labs Bike Light'. Click the shopping cart icon to navigate to "
        "the cart page. Verify both items appear with their names and prices."
    )
    doc.add_paragraph(
        "Expected Result: Cart page lists both items with correct details."
    )

    # --- Section 3: Checkout Tests ---
    doc.add_heading("Checkout Tests", level=1)
    doc.add_paragraph("End-to-end checkout flow validation.")

    # TC-007
    doc.add_heading("TC-007: Complete Checkout Flow", level=2)
    doc.add_paragraph(
        "Log in as standard_user. Add 'Sauce Labs Onesie' to the cart. "
        "Go to the cart page and click 'Checkout'. Fill in 'Test' for first name, "
        "'User' for last name, and '90210' for zip code. Click 'Continue'. "
        "On the overview page verify the item and total are correct. "
        "Click 'Finish'. The confirmation page should say "
        "'Thank you for your order!'"
    )
    doc.add_paragraph(
        "Expected Result: Order confirmation page with thank-you message."
    )

    # TC-008
    doc.add_heading("TC-008: Checkout with Missing Info", level=2)
    doc.add_paragraph(
        "Log in and add an item to the cart. Go to the cart and click Checkout. "
        "Leave the first name field empty and click Continue. "
        "An error should appear: 'Error: First Name is required'."
    )
    doc.add_paragraph(
        "Expected Result: Validation error shown for missing first name."
    )

    path = FIXTURES_DIR / "sample_test_doc_paragraph_steps.docx"
    doc.save(str(path))
    print(f"Created {path}")


# ---------------------------------------------------------------------------
# 3. sample_test_doc_unstructured.docx — Free-form prose, no tables
# ---------------------------------------------------------------------------
def generate_unstructured_doc():
    doc = Document()

    title = doc.add_paragraph("SauceDemo — Exploratory QA Notes")
    title.style = doc.styles["Title"]

    add_metadata(doc, [
        "Prepared by: QA Analyst",
        "Application under test: https://www.saucedemo.com",
        "Date: 2026-03-15",
    ])

    doc.add_paragraph(
        "This document captures test scenarios in an informal, unstructured "
        "manner. The QA agent's LLM parser should be able to extract "
        "actionable test cases from this prose."
    )

    doc.add_heading("Login Scenarios", level=1)

    doc.add_paragraph(
        "First, let's make sure the basic login works. Go to the site and use "
        "standard_user with password secret_sauce — should land on the products page. "
        "Also try an invalid combo like 'foo' / 'bar' and confirm we get an error "
        "message on the page. There's a known locked-out account (locked_out_user / "
        "secret_sauce) that should show a specific lockout message."
    )

    doc.add_heading("Product Browsing", level=1)

    doc.add_paragraph(
        "Once logged in, the products page should list six items. Each item shows "
        "a name, description, price, and an 'Add to cart' button. We should verify "
        "that clicking on a product name opens a detail page with a larger image "
        "and a 'Back to products' link."
    )

    doc.add_paragraph(
        "Sorting is important too. The dropdown at the top right lets you sort by "
        "name A-Z, name Z-A, price low-to-high, and price high-to-low. Pick "
        "price low-to-high and make sure the cheapest item ($7.99 Sauce Labs Onesie) "
        "appears first."
    )

    doc.add_heading("Cart and Checkout", level=1)

    doc.add_paragraph(
        "Add a couple of items to the cart — say the Backpack and the Bike Light. "
        "The cart badge should update to '2'. Open the cart, verify both items are "
        "listed, and then proceed to checkout. Fill in some dummy info (first: Test, "
        "last: User, zip: 12345) and continue to the overview page. Confirm the "
        "item total matches what we'd expect, then finish the order. We should "
        "see a 'Thank you for your order!' confirmation."
    )

    doc.add_paragraph(
        "Edge case: try checking out with an empty cart — the app might still "
        "let you through to the checkout form, but it's worth noting. Also test "
        "checkout validation by leaving the first name blank; we expect an "
        "error 'Error: First Name is required'."
    )

    doc.add_heading("Navigation and Misc", level=1)

    doc.add_paragraph(
        "The hamburger menu (top left) has four options: All Items, About, Logout, "
        "and Reset App State. Verify that 'Logout' returns you to the login page. "
        "'Reset App State' should clear the cart and revert all buttons. 'About' "
        "should navigate to saucelabs.com. 'All Items' should bring you back to "
        "the products page from anywhere in the app."
    )

    doc.add_paragraph(
        "Finally, check the footer links (Twitter, Facebook, LinkedIn) — they "
        "should open the correct Sauce Labs social media pages in new tabs."
    )

    path = FIXTURES_DIR / "sample_test_doc_unstructured.docx"
    doc.save(str(path))
    print(f"Created {path}")


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------
if __name__ == "__main__":
    generate_sample_test_doc()
    generate_paragraph_steps_doc()
    generate_unstructured_doc()
    print("\nAll 3 fixture documents generated successfully.")
